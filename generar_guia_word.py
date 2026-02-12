#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Genera la guía completa del sistema de monitoreo en formato Word (.docx)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ============================================================
# ESTILOS
# ============================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Estilo para código/comandos
if 'Code' not in [s.name for s in doc.styles]:
    code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(10)
    code_font.color.rgb = RGBColor(30, 30, 30)
    code_style.paragraph_format.space_before = Pt(4)
    code_style.paragraph_format.space_after = Pt(4)
    code_style.paragraph_format.left_indent = Cm(1)

# ============================================================
# PORTADA
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
titulo_portada = doc.add_paragraph()
titulo_portada.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = titulo_portada.add_run('Sistema de monitoreo de noticias\nsobre juventud en Colombia')
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0, 51, 102)
run.bold = True

doc.add_paragraph()
subtitulo = doc.add_paragraph()
subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitulo.add_run('Guía completa para el equipo')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
entidad = doc.add_paragraph()
entidad.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = entidad.add_run('Subdirección para la juventud - SDIS\nBogotá, febrero 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_page_break()

# ============================================================
# TABLA DE CONTENIDO (manual)
# ============================================================
doc.add_heading('Contenido', level=1)
contenido = [
    '1. Qué es este sistema y para qué sirve',
    '2. Qué necesita el computador para que funcione',
    '3. Estructura del proyecto (archivos y carpetas)',
    '4. Cómo ejecutar el sistema paso a paso',
    '5. Cómo usar el dashboard (la interfaz visual)',
    '6. Las fuentes de noticias y cómo funcionan',
    '7. Los términos de búsqueda',
    '8. Problemas conocidos y cómo resolverlos',
    '9. Comandos útiles (referencia rápida)',
    '10. Cómo mejorar el sistema',
    '11. Tecnologías recomendadas para un monitoreo de nivel profesional',
    '12. Preguntas frecuentes',
    '13. Glosario de términos técnicos',
]
for item in contenido:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ============================================================
# 1. QUÉ ES ESTE SISTEMA
# ============================================================
doc.add_heading('1. Qué es este sistema y para qué sirve', level=1)

doc.add_paragraph(
    'Este sistema revisa automáticamente las principales páginas de noticias de Colombia '
    'todos los días y recopila las noticias que mencionan temas de juventud (jóvenes, '
    'adolescentes, niños, colegios, etc.).'
)

doc.add_paragraph(
    'Los resultados se guardan en un archivo Excel y se pueden consultar en un dashboard '
    '(tablero visual) con filtros por fecha, ciudad, fuente y tipo de medio.'
)

doc.add_heading('Para qué sirve', level=2)
beneficios = [
    'Tener un panorama diario de lo que se publica en medios sobre juventud en Colombia',
    'No tener que revisar manualmente cada medio todos los días',
    'Identificar tendencias: qué ciudades aparecen más, qué temas son recurrentes',
    'Tener un archivo histórico consultable de noticias sobre juventud',
    'Distinguir entre medios gratuitos y medios con suscripción paga',
]
for b in beneficios:
    doc.add_paragraph(b, style='List Bullet')

doc.add_heading('Qué NO hace este sistema', level=2)
limitaciones = [
    'No lee el contenido completo de artículos detrás de paywall (solo el título y resumen visible)',
    'No analiza el sentimiento ni clasifica las noticias como positivas o negativas',
    'No genera alertas automáticas (solo recopila y muestra)',
    'No garantiza capturar el 100% de las noticias (depende de la estructura de cada sitio web)',
]
for l in limitaciones:
    doc.add_paragraph(l, style='List Bullet')

# ============================================================
# 2. REQUISITOS
# ============================================================
doc.add_heading('2. Qué necesita el computador para que funcione', level=1)

doc.add_heading('Software necesario', level=2)

tabla_req = doc.add_table(rows=4, cols=3)
tabla_req.style = 'Light Shading Accent 1'
tabla_req.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Software', 'Versión mínima', 'Para qué se usa']
for i, h in enumerate(headers):
    tabla_req.rows[0].cells[i].text = h
    for paragraph in tabla_req.rows[0].cells[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True

datos_req = [
    ['Python', '3.11 o superior', 'Ejecutar el scraper y el dashboard'],
    ['pip', 'Viene con Python', 'Instalar las librerías necesarias'],
    ['Git', 'Cualquier versión', 'Subir cambios al repositorio en GitHub'],
]
for i, fila in enumerate(datos_req):
    for j, valor in enumerate(fila):
        tabla_req.rows[i+1].cells[j].text = valor

doc.add_paragraph()
doc.add_heading('Librerías de Python', level=2)
doc.add_paragraph(
    'Todas las librerías necesarias están listadas en el archivo requirements.txt. '
    'Se instalan con un solo comando (ver sección de comandos).'
)

tabla_libs = doc.add_table(rows=8, cols=2)
tabla_libs.style = 'Light Shading Accent 1'
tabla_libs.alignment = WD_TABLE_ALIGNMENT.CENTER

tabla_libs.rows[0].cells[0].text = 'Librería'
tabla_libs.rows[0].cells[1].text = 'Qué hace'
for paragraph in tabla_libs.rows[0].cells[0].paragraphs:
    for run in paragraph.runs:
        run.bold = True
for paragraph in tabla_libs.rows[0].cells[1].paragraphs:
    for run in paragraph.runs:
        run.bold = True

libs = [
    ['requests', 'Hace las peticiones HTTP a los sitios web (como si abriera un navegador)'],
    ['beautifulsoup4 + lxml', 'Lee y entiende el HTML de las páginas web'],
    ['pandas + openpyxl', 'Maneja los datos en tablas y genera el archivo Excel'],
    ['streamlit', 'Crea el dashboard visual (la página web local para consultar)'],
    ['plotly', 'Genera los gráficos interactivos del dashboard'],
    ['feedparser', 'Lee feeds RSS (formato estándar de noticias)'],
    ['python-docx', 'Genera documentos Word (como esta guía)'],
]
for i, (lib, desc) in enumerate(libs):
    tabla_libs.rows[i+1].cells[0].text = lib
    tabla_libs.rows[i+1].cells[1].text = desc

# ============================================================
# 3. ESTRUCTURA DEL PROYECTO
# ============================================================
doc.add_page_break()
doc.add_heading('3. Estructura del proyecto (archivos y carpetas)', level=1)

doc.add_paragraph(
    'El proyecto está organizado de la siguiente manera. Cada archivo tiene una función específica:'
)

tabla_archivos = doc.add_table(rows=12, cols=2)
tabla_archivos.style = 'Light Shading Accent 1'
tabla_archivos.alignment = WD_TABLE_ALIGNMENT.CENTER

tabla_archivos.rows[0].cells[0].text = 'Archivo'
tabla_archivos.rows[0].cells[1].text = 'Qué hace'
for paragraph in tabla_archivos.rows[0].cells[0].paragraphs:
    for run in paragraph.runs:
        run.bold = True
for paragraph in tabla_archivos.rows[0].cells[1].paragraphs:
    for run in paragraph.runs:
        run.bold = True

archivos = [
    ['config.py', 'Configuración: lista de fuentes de noticias, términos de búsqueda y ciudades'],
    ['scraper.py', 'El programa principal que recoge las noticias de cada medio'],
    ['dashboard.py', 'La interfaz visual (dashboard) para consultar las noticias'],
    ['requirements.txt', 'Lista de librerías necesarias para que todo funcione'],
    ['ejecutar_scraper.bat', 'Archivo para ejecutar el scraper con doble clic en Windows'],
    ['data/noticias.xlsx', 'La base de datos con todas las noticias recolectadas'],
    ['.github/workflows/scraper.yml', 'Configuración para ejecución automática diaria en GitHub'],
    ['config_v1.py', 'Copia de respaldo de la configuración original'],
    ['scraper_v1.py', 'Copia de respaldo del scraper original'],
    ['scraper_v2.py', 'Copia de respaldo del scraper antes de las correcciones finales'],
    ['dashboard_v1.py', 'Copia de respaldo del dashboard original'],
]
for i, (archivo, desc) in enumerate(archivos):
    tabla_archivos.rows[i+1].cells[0].text = archivo
    tabla_archivos.rows[i+1].cells[1].text = desc

doc.add_paragraph()
doc.add_paragraph(
    'Las copias de respaldo (v1, v2) existen para poder volver a una versión anterior '
    'si algo falla después de un cambio.'
)

# ============================================================
# 4. CÓMO EJECUTAR
# ============================================================
doc.add_page_break()
doc.add_heading('4. Cómo ejecutar el sistema paso a paso', level=1)

doc.add_heading('Primera vez: instalación', level=2)
doc.add_paragraph('Estos pasos solo se hacen una vez, cuando se configura el proyecto por primera vez:')

doc.add_paragraph('Paso 1: Abrir una terminal (PowerShell o CMD) en la carpeta del proyecto.')
doc.add_paragraph('Paso 2: Instalar las dependencias:', style='List Number')
p = doc.add_paragraph('pip install -r requirements.txt', style='Code')

doc.add_paragraph(
    'Esto descarga e instala todas las librerías necesarias. Solo se hace una vez '
    '(o cuando se agreguen nuevas librerías al archivo requirements.txt).'
)

doc.add_heading('Opción A: ejecutar el scraper manualmente', level=2)
doc.add_paragraph('Se usa cuando se quiere recoger noticias en el momento:')
pasos_manual = [
    'Abrir una terminal en la carpeta del proyecto',
    'Ejecutar el comando: python scraper.py',
    'Esperar a que termine (toma entre 2 y 5 minutos porque visita varias páginas y Google News)',
    'Los resultados se guardan automáticamente en data/noticias.xlsx',
]
for paso in pasos_manual:
    doc.add_paragraph(paso, style='List Number')

doc.add_heading('Opción B: doble clic en Windows', level=2)
doc.add_paragraph(
    'El archivo ejecutar_scraper.bat permite ejecutar el scraper con doble clic. '
    'El resultado queda registrado en la carpeta logs/.'
)
doc.add_paragraph(
    'Nota importante: la ruta dentro de este archivo .bat debe coincidir con la '
    'ubicación real del proyecto en el computador. Si se mueve la carpeta, '
    'se debe actualizar la ruta dentro del archivo.'
)

doc.add_heading('Opción C: ejecución automática con GitHub Actions', level=2)
doc.add_paragraph(
    'Si el proyecto está en GitHub (que ya lo está), el scraper se ejecuta '
    'automáticamente todos los días a las 7:00 AM hora Colombia. '
    'Los resultados se guardan en el repositorio automáticamente.'
)
doc.add_paragraph(
    'Para que esto funcione, el workflow debe estar habilitado en la pestaña '
    '"Actions" del repositorio en GitHub.'
)

doc.add_heading('Ver el dashboard', level=2)
pasos_dashboard = [
    'Abrir una terminal en la carpeta del proyecto',
    'Ejecutar: python -m streamlit run dashboard.py',
    'Se abre automáticamente una página web en el navegador (http://localhost:8501)',
    'Usar los filtros del panel izquierdo para explorar las noticias',
    'Para cerrar el dashboard, presionar Ctrl+C en la terminal',
]
for paso in pasos_dashboard:
    doc.add_paragraph(paso, style='List Number')

# ============================================================
# 5. CÓMO USAR EL DASHBOARD
# ============================================================
doc.add_page_break()
doc.add_heading('5. Cómo usar el dashboard (la interfaz visual)', level=1)

doc.add_paragraph(
    'El dashboard es una página web que se ejecuta localmente (en tu computador). '
    'No necesita internet para funcionar una vez que los datos están descargados.'
)

doc.add_heading('Panel izquierdo: filtros', level=2)

filtros = [
    ['Tipo de fuente', 'Permite ver solo medios gratuitos, solo diarios pagos, o ambos'],
    ['Fuente', 'Seleccionar qué medios específicos mostrar (Pulzo, Infobae, El Tiempo, etc.)'],
    ['Ciudad', 'Filtrar por ciudad mencionada en la noticia'],
    ['Fecha', 'Rango de fechas para acotar el período de consulta'],
    ['Término de búsqueda', 'Filtrar por un término específico de la lista (jóvenes, colegios, etc.)'],
    ['Buscar en títulos', 'Búsqueda libre de texto en los títulos de las noticias'],
    ['Actualizar datos', 'Botón para recargar los datos del Excel (útil después de ejecutar el scraper)'],
]

tabla_filtros = doc.add_table(rows=len(filtros)+1, cols=2)
tabla_filtros.style = 'Light Shading Accent 1'
tabla_filtros.rows[0].cells[0].text = 'Filtro'
tabla_filtros.rows[0].cells[1].text = 'Qué hace'
for paragraph in tabla_filtros.rows[0].cells[0].paragraphs:
    for run in paragraph.runs:
        run.bold = True
for paragraph in tabla_filtros.rows[0].cells[1].paragraphs:
    for run in paragraph.runs:
        run.bold = True
for i, (filtro, desc) in enumerate(filtros):
    tabla_filtros.rows[i+1].cells[0].text = filtro
    tabla_filtros.rows[i+1].cells[1].text = desc

doc.add_paragraph()
doc.add_heading('Panel central: métricas y gráficos', level=2)
metricas = [
    'Total de noticias filtradas',
    'Cantidad de fuentes activas',
    'Cantidad de ciudades mencionadas',
    'Noticias de hoy',
    'Noticias de la última semana',
    'Gráfico de barras: noticias por fuente',
    'Gráfico circular: noticias por ciudad',
    'Gráfico de línea: noticias por día',
    'Gráfico de barras horizontal: términos más frecuentes',
]
for m in metricas:
    doc.add_paragraph(m, style='List Bullet')

doc.add_heading('Tabla de noticias', level=2)
doc.add_paragraph(
    'Debajo de los gráficos está la tabla completa de noticias con columnas de '
    'fecha, título, fuente, tipo (gratuito/diario pago), ciudad, link y resumen. '
    'Se puede descargar en formato CSV con el botón "Descargar CSV".'
)

# ============================================================
# 6. FUENTES DE NOTICIAS
# ============================================================
doc.add_page_break()
doc.add_heading('6. Las fuentes de noticias y cómo funcionan', level=1)

doc.add_paragraph(
    'El sistema monitorea 9 fuentes de noticias colombianas más Google News para búsqueda histórica. '
    'Cada fuente tiene una técnica de extracción diferente porque cada sitio web '
    'está construido de forma distinta.'
)

doc.add_heading('Fuentes gratuitas', level=2)

tabla_fuentes = doc.add_table(rows=8, cols=3)
tabla_fuentes.style = 'Light Shading Accent 1'
tabla_fuentes.rows[0].cells[0].text = 'Fuente'
tabla_fuentes.rows[0].cells[1].text = 'Método de extracción'
tabla_fuentes.rows[0].cells[2].text = 'Notas'
for j in range(3):
    for paragraph in tabla_fuentes.rows[0].cells[j].paragraphs:
        for run in paragraph.runs:
            run.bold = True

fuentes_gratis = [
    ['Blu Radio', 'HTML directo (article tags)', 'Sección nacional'],
    ['Noticias Caracol', 'HTML directo (article tags)', 'Sección Colombia'],
    ['Alerta Bogotá', 'HTML directo (article tags)', 'Noticias de Bogotá por defecto'],
    ['Red+', 'HTML directo (article tags)', 'Página principal'],
    ['Pulzo', 'HTML (a.link-title)', 'Se filtra para evitar autores y números de posición'],
    ['Infobae', 'RSS feed (XML)', 'Trae ~100 artículos recientes de toda Latinoamérica'],
    ['Diario ADN', 'HTML directo (article tags)', 'Página principal'],
]
for i, (fuente, metodo, nota) in enumerate(fuentes_gratis):
    tabla_fuentes.rows[i+1].cells[0].text = fuente
    tabla_fuentes.rows[i+1].cells[1].text = metodo
    tabla_fuentes.rows[i+1].cells[2].text = nota

doc.add_paragraph()
doc.add_heading('Diarios pagos', level=2)

tabla_pagos = doc.add_table(rows=3, cols=3)
tabla_pagos.style = 'Light Shading Accent 1'
tabla_pagos.rows[0].cells[0].text = 'Fuente'
tabla_pagos.rows[0].cells[1].text = 'Método de extracción'
tabla_pagos.rows[0].cells[2].text = 'Notas'
for j in range(3):
    for paragraph in tabla_pagos.rows[0].cells[j].paragraphs:
        for run in paragraph.runs:
            run.bold = True

fuentes_pagas = [
    ['El Tiempo', 'JSON-LD (datos estructurados para SEO)', 'Usa Marfeel (renderiza con JavaScript). El HTML no tiene artículos visibles, pero los datos de SEO sí están disponibles (~16 artículos)'],
    ['El Espectador', 'JSON de Fusion.globalContent', 'Usa Arc Publishing. El JSON embebido en la página tiene ~38 artículos con título y descripción completos'],
]
for i, (fuente, metodo, nota) in enumerate(fuentes_pagas):
    tabla_pagos.rows[i+1].cells[0].text = fuente
    tabla_pagos.rows[i+1].cells[1].text = metodo
    tabla_pagos.rows[i+1].cells[2].text = nota

doc.add_paragraph()
doc.add_heading('Búsqueda histórica (Google News)', level=2)
doc.add_paragraph(
    'Además de las fuentes directas, el sistema busca en Google News RSS para '
    'recuperar noticias de los últimos 60 días. Esto permite capturar noticias '
    'que no se encontraron en las páginas del día porque ya no están en portada.'
)
doc.add_paragraph(
    'Google News agrega la fuente original de cada noticia, así que en el Excel '
    'pueden aparecer fuentes adicionales a las 9 configuradas (como RCN, Semana, etc.).'
)

doc.add_heading('Por qué cada fuente necesita un método diferente', level=2)
doc.add_paragraph(
    'Cada sitio web está construido con tecnología diferente. Algunos muestran las noticias '
    'directamente en el HTML (como Blu Radio o Caracol), pero otros usan JavaScript para '
    'cargar el contenido después (como El Tiempo con Marfeel o Infobae con Arc Publishing). '
    'Cuando el contenido se carga con JavaScript, el scraper no puede "verlo" porque no ejecuta '
    'JavaScript como un navegador. En esos casos se usan fuentes alternativas de datos:'
)
alternativas = [
    'JSON-LD: datos que los sitios ponen en su HTML para que Google los indexe (El Tiempo)',
    'Fusion JSON: datos embebidos por el framework Arc Publishing antes del renderizado (El Espectador)',
    'RSS: formato estándar de noticias que muchos medios ofrecen como feed (Infobae)',
]
for a in alternativas:
    doc.add_paragraph(a, style='List Bullet')

# ============================================================
# 7. TÉRMINOS DE BÚSQUEDA
# ============================================================
doc.add_page_break()
doc.add_heading('7. Los términos de búsqueda', level=1)

doc.add_paragraph(
    'El sistema solo guarda noticias que contengan al menos uno de estos términos '
    'en el título o resumen. Este es el filtro principal que asegura que las noticias '
    'sean relevantes para temas de juventud.'
)

doc.add_heading('Lista actual de términos', level=2)
terminos = [
    'juventud', 'jóvenes', 'jovenes', 'adolescentes', 'adolescencia',
    'menor de edad', 'menores de edad', 'pandillas', 'idipron',
    'plataformas juveniles', 'primera infancia', 'niños', 'niñas',
    'infancia', 'juvenil', 'juveniles', 'estudiantes', 'colegios', 'escolar'
]
for t in terminos:
    doc.add_paragraph(t, style='List Bullet')

doc.add_heading('Cómo agregar o quitar términos', level=2)
doc.add_paragraph(
    'Se edita el archivo config.py, específicamente la lista TERMINOS_JUVENTUD. '
    'Cada término va entre comillas y separado por coma. Ejemplo:'
)
p = doc.add_paragraph('"nuevo término",', style='Code')

doc.add_paragraph(
    'Tener en cuenta que entre más genérico sea el término, más noticias no relevantes '
    'puede traer. Por ejemplo, "estudiantes" puede traer noticias sobre estudiantes '
    'universitarios que no son del interés del equipo.'
)

# ============================================================
# 8. PROBLEMAS CONOCIDOS
# ============================================================
doc.add_heading('8. Problemas conocidos y cómo resolverlos', level=1)

doc.add_heading('Títulos o resúmenes incorrectos', level=2)
doc.add_paragraph(
    'Algunos títulos pueden salir mal porque los sitios web cambian su estructura HTML '
    'sin previo aviso. Cuando un medio hace un rediseño, el scraper puede empezar a '
    'capturar texto de navegación, publicidad o widgets en lugar del título real.'
)
doc.add_paragraph('Qué hacer:', style='List Bullet')
doc.add_paragraph(
    'Si se detectan muchos errores en una fuente específica, se debe revisar la estructura '
    'HTML actual de ese sitio y ajustar la función correspondiente en scraper.py. '
    'Cada medio tiene su propia función (scrape_bluradio, scrape_pulzo, etc.).'
)

doc.add_heading('El Tiempo muestra 0 noticias', level=2)
doc.add_paragraph(
    'Es normal. El Tiempo tiene ~16 artículos disponibles en su JSON-LD, y muy pocos '
    'tratan temas de juventud en un día cualquiera. Si en un día no hay noticias de juventud '
    'en El Tiempo, el scraper reporta 0. No es un error, es que no hay noticias relevantes ese día.'
)

doc.add_heading('El scraper tarda mucho', level=2)
doc.add_paragraph(
    'Con la búsqueda histórica de Google News, el scraper puede tardar entre 3 y 5 minutos. '
    'Esto es normal porque busca 8 términos diferentes en Google News con una pausa de 1 segundo '
    'entre cada uno para no ser bloqueado.'
)

doc.add_heading('Noticias duplicadas', level=2)
doc.add_paragraph(
    'El sistema usa la URL como identificador único. Si la misma noticia aparece con URLs '
    'diferentes (por ejemplo, una de Google News y otra del medio directo), puede aparecer '
    'dos veces. Esto es un caso marginal que raramente ocurre.'
)

doc.add_heading('Noticias de otros países', level=2)
doc.add_paragraph(
    'El RSS de Infobae y Google News pueden traer noticias de otros países latinoamericanos. '
    'El filtro de términos de juventud aplica igual, pero la ciudad puede quedar como '
    '"Sin identificar" si no menciona una ciudad colombiana.'
)

# ============================================================
# 9. COMANDOS ÚTILES
# ============================================================
doc.add_page_break()
doc.add_heading('9. Comandos útiles (referencia rápida)', level=1)

doc.add_paragraph(
    'Todos estos comandos se ejecutan en una terminal (PowerShell o CMD) '
    'ubicada en la carpeta del proyecto.'
)

tabla_cmd = doc.add_table(rows=11, cols=2)
tabla_cmd.style = 'Light Shading Accent 1'
tabla_cmd.rows[0].cells[0].text = 'Comando'
tabla_cmd.rows[0].cells[1].text = 'Qué hace'
for j in range(2):
    for paragraph in tabla_cmd.rows[0].cells[j].paragraphs:
        for run in paragraph.runs:
            run.bold = True

comandos = [
    ['pip install -r requirements.txt', 'Instala todas las librerías necesarias (solo la primera vez)'],
    ['python scraper.py', 'Ejecuta el scraper y recoge noticias de todas las fuentes'],
    ['python -m streamlit run dashboard.py', 'Abre el dashboard visual en el navegador'],
    ['git status', 'Muestra qué archivos han cambiado desde el último commit'],
    ['git add .', 'Prepara todos los archivos modificados para subir'],
    ['git commit -m "descripción"', 'Guarda los cambios con una descripción'],
    ['git push origin main', 'Sube los cambios a GitHub'],
    ['git pull origin main', 'Descarga los cambios más recientes de GitHub'],
    ['Ctrl + C', 'Detiene el dashboard o cualquier proceso en la terminal'],
    ['cls (Windows) / clear (Mac/Linux)', 'Limpia la pantalla de la terminal'],
]
for i, (cmd, desc) in enumerate(comandos):
    tabla_cmd.rows[i+1].cells[0].text = cmd
    tabla_cmd.rows[i+1].cells[1].text = desc

# ============================================================
# 10. CÓMO MEJORAR EL SISTEMA
# ============================================================
doc.add_page_break()
doc.add_heading('10. Cómo mejorar el sistema', level=1)

doc.add_paragraph(
    'El sistema actual cumple su función básica de monitoreo, pero tiene espacio '
    'para mejorar significativamente. Estas son las mejoras posibles ordenadas '
    'de menor a mayor esfuerzo:'
)

doc.add_heading('Mejoras rápidas (se pueden hacer en 1-2 horas)', level=2)

mejoras_rapidas = [
    ('Entrar al artículo para extraer el título real',
     'En lugar de tomar el título de la tarjeta en la página de sección, '
     'el scraper entraría a cada artículo y extraería el <h1> principal. '
     'Esto da títulos más limpios y permite extraer también el texto completo del artículo. '
     'Desventaja: el scraper se vuelve más lento porque hace más peticiones HTTP.'),

    ('Agregar más fuentes de noticias',
     'Se pueden agregar fuentes como Semana, RCN Noticias, W Radio, La FM, '
     'El Colombiano (Medellín), El País (Cali), El Heraldo (Barranquilla). '
     'Cada fuente requiere escribir una nueva función de scraping adaptada a su HTML.'),

    ('Filtrar noticias internacionales',
     'Agregar un filtro que descarte noticias que no mencionen Colombia o ciudades colombianas, '
     'para evitar noticias de Argentina, Ecuador, etc. que llegan por Infobae y Google News.'),

    ('Mejorar la detección de ciudad',
     'El sistema actual busca nombres de ciudades en el texto. Se podría usar '
     'la sección/categoría del medio (que a veces indica la ciudad) como fuente adicional.'),
]
for titulo_m, desc_m in mejoras_rapidas:
    doc.add_heading(titulo_m, level=3)
    doc.add_paragraph(desc_m)

doc.add_heading('Mejoras medianas (requieren 1-2 días)', level=2)

mejoras_medianas = [
    ('Clasificación automática por tema',
     'Usar categorías como "seguridad", "educación", "salud", "política", "deportes" '
     'para clasificar automáticamente cada noticia. Se puede hacer con palabras clave '
     'o con un modelo de lenguaje (IA).'),

    ('Alertas por correo electrónico',
     'Enviar un correo diario al equipo con las noticias nuevas encontradas. '
     'Se puede hacer con la librería smtplib de Python o con servicios como SendGrid.'),

    ('Análisis de sentimiento',
     'Determinar si cada noticia es positiva, negativa o neutral. '
     'Se puede hacer con modelos como VADER (simple) o con la API de Claude/GPT (más preciso).'),

    ('Programar ejecución en el computador',
     'Usar el Programador de Tareas de Windows para ejecutar el scraper '
     'automáticamente cada día sin necesidad de GitHub Actions. '
     'Útil si se quiere mantener todo local.'),
]
for titulo_m, desc_m in mejoras_medianas:
    doc.add_heading(titulo_m, level=3)
    doc.add_paragraph(desc_m)

doc.add_heading('Mejoras grandes (requieren 1+ semanas)', level=2)

mejoras_grandes = [
    ('Desplegar el dashboard en la nube',
     'En lugar de ejecutar el dashboard localmente, publicarlo en un servidor '
     'para que todo el equipo pueda accederlo desde cualquier computador con un enlace. '
     'Opciones: Streamlit Community Cloud (gratis), Railway, Render, o un servidor de SDIS.'),

    ('Base de datos en vez de Excel',
     'Reemplazar el Excel por una base de datos (PostgreSQL o SQLite) para manejar '
     'miles de noticias sin problemas de rendimiento. El Excel funciona bien hasta '
     '~5.000 noticias, después se vuelve lento.'),

    ('Resumen automático con IA',
     'Usar la API de Claude o GPT para generar resúmenes concisos de cada noticia, '
     'clasificarla por tema, y extraer entidades mencionadas (personas, instituciones, lugares).'),
]
for titulo_m, desc_m in mejoras_grandes:
    doc.add_heading(titulo_m, level=3)
    doc.add_paragraph(desc_m)

# ============================================================
# 11. TOP TECNOLOGÍAS
# ============================================================
doc.add_page_break()
doc.add_heading('11. Tecnologías recomendadas para un monitoreo de nivel profesional', level=1)

doc.add_paragraph(
    'Si se quisiera llevar este sistema al siguiente nivel, estas son las '
    'tecnologías y herramientas que usan los equipos de monitoreo de medios profesionales. '
    'Se organizan por categoría:'
)

doc.add_heading('Scraping y recolección de datos', level=2)

tech_scraping = doc.add_table(rows=6, cols=3)
tech_scraping.style = 'Light Shading Accent 1'
tech_scraping.rows[0].cells[0].text = 'Tecnología'
tech_scraping.rows[0].cells[1].text = 'Qué hace'
tech_scraping.rows[0].cells[2].text = 'Cuándo usarla'
for j in range(3):
    for paragraph in tech_scraping.rows[0].cells[j].paragraphs:
        for run in paragraph.runs:
            run.bold = True

scraping_tech = [
    ['Playwright', 'Abre un navegador real automatizado que ejecuta JavaScript', 'Cuando el sitio carga contenido con JavaScript (como El Tiempo con Marfeel). Es más moderno y rápido que Selenium'],
    ['Scrapy', 'Framework profesional de scraping con manejo de concurrencia', 'Cuando se necesita scrapear muchas páginas rápidamente y de forma eficiente. Tiene sistema de pipelines y middlewares'],
    ['News API (newsapi.org)', 'API comercial que agrega noticias de miles de fuentes', 'Si el presupuesto lo permite (~$449/mes plan Business). Da acceso a noticias históricas con filtros avanzados'],
    ['GDELT Project', 'Base de datos global gratuita de noticias', 'Para análisis a gran escala de medios. Tiene API gratuita con cobertura de Colombia'],
    ['Common Crawl', 'Archivo masivo de páginas web (petabytes)', 'Para investigación a gran escala. Requiere conocimiento de procesamiento distribuido (Spark/Hadoop)'],
]
for i, (tech, desc, cuando) in enumerate(scraping_tech):
    tech_scraping.rows[i+1].cells[0].text = tech
    tech_scraping.rows[i+1].cells[1].text = desc
    tech_scraping.rows[i+1].cells[2].text = cuando

doc.add_paragraph()
doc.add_heading('Análisis e inteligencia artificial', level=2)

tech_ia = doc.add_table(rows=6, cols=3)
tech_ia.style = 'Light Shading Accent 1'
tech_ia.rows[0].cells[0].text = 'Tecnología'
tech_ia.rows[0].cells[1].text = 'Qué hace'
tech_ia.rows[0].cells[2].text = 'Nivel de dificultad'
for j in range(3):
    for paragraph in tech_ia.rows[0].cells[j].paragraphs:
        for run in paragraph.runs:
            run.bold = True

ia_tech = [
    ['API de Claude (Anthropic)', 'Análisis de texto, clasificación, resúmenes, extracción de entidades', 'Medio. Requiere API key y créditos (~$3/1M tokens). Muy preciso para español'],
    ['spaCy (modelo es_core_news_lg)', 'Procesamiento de lenguaje natural: detecta personas, lugares, organizaciones', 'Bajo-Medio. Gratis, se ejecuta local. Bueno para extracción de entidades'],
    ['Hugging Face Transformers', 'Modelos de IA para clasificación de texto, sentimiento, resúmenes', 'Medio-Alto. Gratis pero requiere GPU para modelos grandes. Muchos modelos en español'],
    ['VADER / TextBlob', 'Análisis de sentimiento simple (positivo/negativo/neutral)', 'Bajo. Gratis, rápido, pero menos preciso para español'],
    ['LangChain + LLM', 'Orquestación de modelos de IA para flujos complejos', 'Alto. Permite crear pipelines: scraping → clasificación → resumen → alerta'],
]
for i, (tech, desc, nivel) in enumerate(ia_tech):
    tech_ia.rows[i+1].cells[0].text = tech
    tech_ia.rows[i+1].cells[1].text = desc
    tech_ia.rows[i+1].cells[2].text = nivel

doc.add_paragraph()
doc.add_heading('Almacenamiento y base de datos', level=2)

tech_db = doc.add_table(rows=5, cols=3)
tech_db.style = 'Light Shading Accent 1'
tech_db.rows[0].cells[0].text = 'Tecnología'
tech_db.rows[0].cells[1].text = 'Qué hace'
tech_db.rows[0].cells[2].text = 'Recomendación'
for j in range(3):
    for paragraph in tech_db.rows[0].cells[j].paragraphs:
        for run in paragraph.runs:
            run.bold = True

db_tech = [
    ['SQLite', 'Base de datos local en un archivo (como Excel pero más rápido)', 'Ideal como siguiente paso. No requiere servidor, solo un archivo .db. Soporta cientos de miles de registros'],
    ['PostgreSQL', 'Base de datos profesional con búsqueda full-text', 'Para cuando el proyecto crezca a múltiples usuarios y se despliegue en servidor'],
    ['Supabase', 'PostgreSQL en la nube con API automática (gratis hasta 500MB)', 'Si se quiere acceder a los datos desde múltiples aplicaciones sin montar servidor'],
    ['Elasticsearch', 'Motor de búsqueda de texto completo, muy rápido', 'Para búsquedas avanzadas en miles de noticias. Requiere más infraestructura'],
]
for i, (tech, desc, rec) in enumerate(db_tech):
    tech_db.rows[i+1].cells[0].text = tech
    tech_db.rows[i+1].cells[1].text = desc
    tech_db.rows[i+1].cells[2].text = rec

doc.add_paragraph()
doc.add_heading('Visualización y dashboard', level=2)

tech_viz = doc.add_table(rows=5, cols=3)
tech_viz.style = 'Light Shading Accent 1'
tech_viz.rows[0].cells[0].text = 'Tecnología'
tech_viz.rows[0].cells[1].text = 'Qué hace'
tech_viz.rows[0].cells[2].text = 'Recomendación'
for j in range(3):
    for paragraph in tech_viz.rows[0].cells[j].paragraphs:
        for run in paragraph.runs:
            run.bold = True

viz_tech = [
    ['Streamlit (actual)', 'Dashboard rápido en Python, ideal para prototipos', 'Lo que se usa actualmente. Perfecto para el equipo actual. Se puede desplegar gratis en Streamlit Cloud'],
    ['Power BI', 'Herramienta de Microsoft para dashboards empresariales', 'Si SDIS ya tiene licencia de Microsoft 365. Más familiar para equipos no técnicos'],
    ['Metabase', 'Dashboard open source que se conecta a bases de datos', 'Si se migra a PostgreSQL. Interfaz intuitiva sin necesidad de programar'],
    ['Grafana', 'Dashboards de monitoreo en tiempo real', 'Si se necesitan alertas en tiempo real y métricas de tendencia'],
]
for i, (tech, desc, rec) in enumerate(viz_tech):
    tech_viz.rows[i+1].cells[0].text = tech
    tech_viz.rows[i+1].cells[1].text = desc
    tech_viz.rows[i+1].cells[2].text = rec

doc.add_paragraph()
doc.add_heading('Automatización y despliegue', level=2)

tech_deploy = doc.add_table(rows=5, cols=3)
tech_deploy.style = 'Light Shading Accent 1'
tech_deploy.rows[0].cells[0].text = 'Tecnología'
tech_deploy.rows[0].cells[1].text = 'Qué hace'
tech_deploy.rows[0].cells[2].text = 'Recomendación'
for j in range(3):
    for paragraph in tech_deploy.rows[0].cells[j].paragraphs:
        for run in paragraph.runs:
            run.bold = True

deploy_tech = [
    ['GitHub Actions (actual)', 'Ejecución automática diaria del scraper', 'Lo que se usa actualmente. Gratis para repositorios públicos'],
    ['Streamlit Community Cloud', 'Hosting gratuito para dashboards Streamlit', 'La forma más fácil de compartir el dashboard con el equipo sin configurar servidores'],
    ['Railway / Render', 'Plataformas de despliegue con plan gratuito', 'Para desplegar tanto el scraper como el dashboard en la nube'],
    ['Docker', 'Empaqueta todo el proyecto para que funcione igual en cualquier máquina', 'Útil si el proyecto se mueve a un servidor de SDIS o se comparte con otros equipos'],
]
for i, (tech, desc, rec) in enumerate(deploy_tech):
    tech_deploy.rows[i+1].cells[0].text = tech
    tech_deploy.rows[i+1].cells[1].text = desc
    tech_deploy.rows[i+1].cells[2].text = rec

doc.add_paragraph()
doc.add_heading('Ruta recomendada de mejora (paso a paso)', level=2)
doc.add_paragraph(
    'No se necesita implementar todo a la vez. Esta es la secuencia recomendada:'
)
ruta = [
    'Desplegar el dashboard en Streamlit Cloud (gratis, 30 minutos) para que todo el equipo pueda verlo',
    'Agregar clasificación por tema usando palabras clave (1 día de trabajo)',
    'Agregar más fuentes colombianas: Semana, RCN, El Colombiano (1-2 días)',
    'Migrar de Excel a SQLite cuando se superen las 5.000 noticias (medio día)',
    'Agregar análisis con API de Claude para resúmenes y clasificación automática (2-3 días)',
    'Agregar alertas por correo electrónico para noticias urgentes (1 día)',
]
for i, paso in enumerate(ruta):
    doc.add_paragraph(f'Fase {i+1}: {paso}', style='List Number')

# ============================================================
# 12. PREGUNTAS FRECUENTES
# ============================================================
doc.add_page_break()
doc.add_heading('12. Preguntas frecuentes', level=1)

faqs = [
    ('¿El scraper puede ser bloqueado por los sitios web?',
     'Sí, es posible. El sistema incluye headers que simulan un navegador normal '
     'y hace pausas de 1 segundo entre peticiones para no saturar los servidores. '
     'Si un sitio bloquea las peticiones, el scraper simplemente reporta 0 noticias '
     'para esa fuente sin afectar las demás.'),

    ('¿Qué pasa si muevo la carpeta del proyecto a otra ubicación?',
     'El proyecto funciona desde cualquier carpeta. Solo se debe actualizar la ruta '
     'dentro del archivo ejecutar_scraper.bat si se usa esa opción de ejecución.'),

    ('¿Se pierden las noticias anteriores cuando se ejecuta el scraper?',
     'No. El scraper siempre agrega noticias nuevas al Excel existente. Nunca borra '
     'las anteriores. Los duplicados se detectan por URL.'),

    ('¿Puedo editar el Excel manualmente?',
     'Sí, se puede abrir en Excel y editar, agregar o eliminar filas. '
     'El scraper respetará los cambios la próxima vez que se ejecute.'),

    ('¿Qué significa "diario pago" en la columna tipo?',
     'Indica que la noticia viene de un medio con suscripción (El Tiempo, El Espectador). '
     'El título y resumen son visibles, pero para leer el artículo completo '
     'puede ser necesario tener suscripción al medio.'),

    ('¿Por qué algunas noticias no tienen resumen?',
     'Depende de la fuente. Algunos feeds RSS y listados de artículos solo incluyen '
     'el título sin descripción. En esos casos, el resumen queda vacío o repite el título.'),

    ('¿Puedo ejecutar el dashboard sin internet?',
     'Sí. El dashboard lee los datos del Excel local. Solo se necesita internet '
     'para ejecutar el scraper (que visita las páginas web).'),

    ('¿Cómo agrego una nueva fuente de noticias?',
     'Se necesitan tres pasos: (1) Agregar la fuente en config.py con su URL y tipo, '
     '(2) Crear una nueva función de scraping en scraper.py adaptada a la estructura HTML '
     'del sitio, (3) Agregar la función a la lista de scrapers en ejecutar_scraping(). '
     'Requiere conocimiento básico de Python y HTML.'),
]

for pregunta, respuesta in faqs:
    doc.add_heading(pregunta, level=3)
    doc.add_paragraph(respuesta)

# ============================================================
# 13. GLOSARIO
# ============================================================
doc.add_page_break()
doc.add_heading('13. Glosario de términos técnicos', level=1)

doc.add_paragraph(
    'Estos son los términos técnicos que aparecen en esta guía y en el proyecto, '
    'explicados de forma simple:'
)

glosario = [
    ('API', 'Una forma de que dos programas se comuniquen entre sí. Es como un "mesero" que lleva pedidos de un programa a otro y trae la respuesta.'),
    ('Commit', 'Un punto de guardado en Git. Como guardar una partida en un videojuego: se puede volver a ese punto si algo sale mal.'),
    ('CSS', 'El código que define cómo se ve una página web (colores, tamaños, posiciones). El scraper a veces busca elementos por sus clases CSS.'),
    ('Dashboard', 'Un tablero visual con gráficos y filtros para consultar datos de forma interactiva.'),
    ('Feed RSS', 'Un formato estándar que usan los sitios de noticias para compartir sus artículos de forma estructurada (como un catálogo digital de noticias).'),
    ('Framework', 'Un conjunto de herramientas y reglas prehechas que facilitan construir software. Streamlit es un framework para dashboards, Scrapy para scraping.'),
    ('Git', 'Sistema de control de versiones. Registra todos los cambios hechos al código y permite volver a versiones anteriores.'),
    ('GitHub', 'Un sitio web donde se almacenan proyectos que usan Git. Es como Google Drive pero para código.'),
    ('GitHub Actions', 'Un servicio de GitHub que ejecuta tareas automáticas (como correr el scraper todos los días).'),
    ('HTML', 'El código que define la estructura de una página web. El scraper lee el HTML para encontrar las noticias.'),
    ('JavaScript', 'Un lenguaje de programación que se ejecuta en el navegador. Algunos sitios cargan su contenido con JavaScript, lo que dificulta el scraping.'),
    ('JSON', 'Un formato de datos ligero (como una tabla pero en texto). Muchos sitios embeben datos de artículos en formato JSON.'),
    ('JSON-LD', 'Un tipo especial de JSON que los sitios web usan para describir su contenido a Google y otros buscadores.'),
    ('Paywall', 'Muro de pago. Cuando un sitio web requiere suscripción para leer el contenido completo.'),
    ('Pipeline', 'Una secuencia de pasos automatizados que procesan datos: recoger → filtrar → clasificar → guardar → alertar.'),
    ('Push', 'Subir los cambios locales a GitHub (el repositorio remoto).'),
    ('Pull', 'Descargar los cambios de GitHub al computador local.'),
    ('Scraper', 'Un programa que visita páginas web automáticamente y extrae información de ellas.'),
    ('SPA (Single Page Application)', 'Un tipo de sitio web que carga todo con JavaScript. Infobae es un ejemplo.'),
    ('Streamlit', 'Una librería de Python para crear dashboards web rápidamente sin necesidad de saber desarrollo web.'),
    ('Terminal', 'La ventana negra (o azul) donde se escriben comandos. En Windows se llama PowerShell o CMD.'),
]

tabla_glosario = doc.add_table(rows=len(glosario)+1, cols=2)
tabla_glosario.style = 'Light Shading Accent 1'
tabla_glosario.rows[0].cells[0].text = 'Término'
tabla_glosario.rows[0].cells[1].text = 'Significado'
for j in range(2):
    for paragraph in tabla_glosario.rows[0].cells[j].paragraphs:
        for run in paragraph.runs:
            run.bold = True
for i, (termino, significado) in enumerate(glosario):
    tabla_glosario.rows[i+1].cells[0].text = termino
    tabla_glosario.rows[i+1].cells[1].text = significado

# ============================================================
# PIE DE DOCUMENTO
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
pie = doc.add_paragraph()
pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = pie.add_run(
    'Documento generado el 12 de febrero de 2026\n'
    'Sistema construido con la asistencia de Claude (Anthropic)\n'
    'Subdirección para la juventud - SDIS - Bogotá'
)
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(128, 128, 128)

# ============================================================
# GUARDAR
# ============================================================
ruta_salida = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           'Guia_sistema_monitoreo_noticias_juventud.docx')
doc.save(ruta_salida)
print(f"Documento guardado en: {ruta_salida}")
